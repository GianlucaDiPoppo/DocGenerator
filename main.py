import customtkinter as ctk
from tkinter import filedialog, messagebox
import pandas as pd
import os, json, requests, webbrowser
from docx import Document

# --- CONFIGURAZIONE ---
VERSIONE_ATTUALE = "v1.0.0"
REPO_GITHUB = "GianlucaDiPoppo/DocGenerator" 
CONFIG_FILE = "config.json"

def carica_config():
    default = {
        "colonne": {"tipo_b": 1, "titolo_d": 3, "numero_e": 4, "p": 15, "q": 16},
        "frasi": {
            "ftr": "All’interno della transazione vengono utilizzati i seguenti FTR:",
            "det": "All’interno della transazione vengono utilizzati i seguenti DET:",
            "ret_dati": "I RET presenti all'interno dell'entità sono:",
            "det_dati": "I DET che rappresentano l'entità sono:"
        },
        "spazi": {"tra_blocchi": 7, "tra_frasi": 1}
    }
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return {**default, **json.load(f)}
    return default

# --- LOGICA CORE ---
def genera_word(percorso_excel, foglio, config):
    try:
        df = pd.read_excel(percorso_excel, sheet_name=foglio, header=None)
        doc = Document()
        c = config
        
        for index, row in df.iterrows():
            tipo = str(row[c['colonne']['tipo_b']]).strip().upper()
            titolo = str(row[c['colonne']['titolo_d']]).strip()
            num = str(row[c['colonne']['numero_e']]).strip()
            if titolo == 'nan' or not titolo: continue
            
            # Titoli automatici (Word gestisce i numeri)
            if tipo == "DATI":
                doc.add_heading(f"{titolo} [{num}]", level=4)
                testi = [(c['frasi']['ret_dati'], row[c['colonne']['p']]), (c['frasi']['det_dati'], row[c['colonne']['q']])]
            else:
                doc.add_heading(f"{titolo} [{num}]", level=3)
                doc.add_heading("Soluzione", level=4)
                testi = [(c['frasi']['ftr'], row[c['colonne']['p']]), (c['frasi']['det'], row[c['colonne']['q']])]

            for i, (intro, contenuto) in enumerate(testi):
                p = doc.add_paragraph()
                p.add_run(intro).italic = True
                if str(contenuto) != 'nan':
                    voci = [v.strip() for v in str(contenuto).replace('\n','|').replace(';','|').split('|') if v.strip()]
                    for v in voci: doc.add_paragraph(v, style='List Bullet')
                if i < len(testi)-1:
                    for _ in range(c['spazi']['tra_frasi']): doc.add_paragraph()

            if index < len(df)-1:
                for _ in range(c['spazi']['tra_blocchi']): doc.add_paragraph()

        out = percorso_excel.replace(".xlsx", "_Documento.docx")
        doc.save(out)
        return True, out
    except Exception as e: return False, str(e)

# --- INTERFACCIA ---
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(f"Doc Generator - {VERSIONE_ATTUALE}")
        self.geometry("600x500")
        self.config_data = carica_config()
        self.tabview = ctk.CTkTabview(self)
        self.tabview.pack(fill="both", expand=True, padx=10, pady=10)
        self.tab_gen = self.tabview.add("Generatore")
        self.tab_set = self.tabview.add("Impostazioni")
        self.init_gen(); self.init_set()
        self.after(1000, self.check_updates)

    def init_gen(self):
        ctk.CTkButton(self.tab_gen, text="Seleziona Excel", command=self.carica).pack(pady=20)
        self.path_lbl = ctk.CTkLabel(self.tab_gen, text="Nessun file selezionato")
        self.path_lbl.pack()
        self.sheet_ent = ctk.CTkEntry(self.tab_gen, placeholder_text="Nome Foglio")
        self.sheet_ent.insert(0, "ConteggioFunzioni")
        self.sheet_ent.pack(pady=10)
        ctk.CTkButton(self.tab_gen, text="AVVIA", fg_color="green", command=self.run).pack(pady=20)

    def init_set(self):
        # aggiungere i campi per modificare config.json
        ctk.CTkLabel(self.tab_set, text="Modifica config.json per cambiare colonne/testi").pack(pady=50)

    def carica(self):
        p = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx")])
        if p: self.path_lbl.configure(text=p)

    def run(self):
        ok, msg = genera_word(self.path_lbl.cget("text"), self.sheet_ent.get(), self.config_data)
        messagebox.showinfo("Esito", f"Creato: {msg}" if ok else f"Errore: {msg}")

    def check_updates(self):
        try:
            r = requests.get(f"https://api.github.com/repos/{REPO_GITHUB}/releases/latest", timeout=3)
            if r.status_code == 200:
                ultima = r.json()["tag_name"]
                if ultima != VERSIONE_ATTUALE:
                    if messagebox.askyesno("Update", f"Nuova versione {ultima} disponibile. Scaricare?"):
                        webbrowser.open(r.json()["html_url"])
        except: pass

if __name__ == "__main__": App().mainloop()